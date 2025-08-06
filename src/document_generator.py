# src/document_generator.py
# Enhanced document generator for ISO13485 compliance documents

import json
from typing import Dict, Optional, List, Any
from io import BytesIO
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn
import streamlit as st
from .ai_integration import DualAPIManager

class ISO13485DocumentGenerator:
    """Enhanced document generator for comprehensive ISO13485 document creation."""
    
    def __init__(self, ai_manager: Optional[DualAPIManager] = None):
        self.ai_manager = ai_manager
        self.document_templates = self._load_document_templates()
        
        # Document formatting standards
        self.formatting = {
            "title_font": "Arial",
            "title_size": 16,
            "header_font": "Arial", 
            "header_size": 12,
            "body_font": "Arial",
            "body_size": 11,
            "colors": {
                "primary": RGBColor(0, 70, 107),  # Dark blue
                "secondary": RGBColor(59, 130, 246),  # Light blue
                "success": RGBColor(16, 185, 129),  # Green
                "warning": RGBColor(245, 158, 11),  # Orange
                "error": RGBColor(239, 68, 68)   # Red
            }
        }
    
    def _load_document_templates(self) -> Dict[str, Dict]:
        """Load document templates and structures."""
        return {
            "CAPA Request Form": {
                "sections": [
                    "Document Header",
                    "CAPA Identification",
                    "Problem Description", 
                    "Immediate Actions",
                    "Root Cause Analysis",
                    "Corrective Actions",
                    "Preventive Actions",
                    "Effectiveness Verification",
                    "Approval and Sign-off"
                ],
                "required_fields": [
                    "capa_number", "product_name", "issue_description", 
                    "root_cause", "corrective_action", "preventive_action"
                ]
            },
            "Nonconformance Report": {
                "sections": [
                    "Document Header",
                    "Nonconformance Identification",
                    "Description of Nonconformance",
                    "Impact Assessment",
                    "Disposition",
                    "Investigation",
                    "Corrective Action Required",
                    "Closure"
                ],
                "required_fields": [
                    "ncr_number", "product_name", "nonconformance_description",
                    "disposition", "investigation_findings"
                ]
            },
            "Internal Audit Checklist": {
                "sections": [
                    "Audit Information",
                    "Audit Scope",
                    "ISO13485 Requirements Checklist",
                    "Findings",
                    "Nonconformities",
                    "Recommendations",
                    "Conclusion"
                ],
                "required_fields": [
                    "audit_number", "audit_date", "audit_scope", "auditor_name"
                ]
            },
            "Risk Assessment Form": {
                "sections": [
                    "Risk Assessment Information",
                    "Hazard Identification",
                    "Risk Analysis",
                    "Risk Evaluation", 
                    "Risk Control Measures",
                    "Residual Risk Assessment",
                    "Risk Management Report"
                ],
                "required_fields": [
                    "product_name", "hazard_identification", "risk_analysis",
                    "risk_control_measures"
                ]
            },
            "Management Review Agenda": {
                "sections": [
                    "Meeting Information",
                    "Review of Previous Management Review",
                    "Quality Policy and Objectives Review",
                    "Quality System Performance",
                    "Customer Feedback",
                    "Audit Results", 
                    "CAPA Status",
                    "Resource Requirements",
                    "Improvement Opportunities"
                ],
                "required_fields": [
                    "review_date", "attendees", "review_period"
                ]
            },
            "Supplier Audit Report": {
                "sections": [
                    "Audit Information",
                    "Supplier Information", 
                    "Audit Scope and Criteria",
                    "Audit Findings",
                    "Nonconformities",
                    "Corrective Actions Required",
                    "Supplier Qualification Status",
                    "Recommendations"
                ],
                "required_fields": [
                    "supplier_name", "audit_date", "audit_scope", "audit_findings"
                ]
            }
        }
    
    def generate_document(
        self, 
        document_type: str, 
        data: Dict[str, Any],
        model: str = "auto"
    ) -> Dict[str, Any]:
        """Generate a complete ISO13485 document using AI."""
        
        if document_type not in self.document_templates:
            raise ValueError(f"Document type '{document_type}' not supported")
        
        template = self.document_templates[document_type]
        
        # Validate required fields
        missing_fields = []
        for field in template["required_fields"]:
            if not data.get(field):
                missing_fields.append(field)
        
        if missing_fields and self.ai_manager:
            # Use AI to suggest content for missing fields
            data = self._ai_fill_missing_fields(document_type, data, missing_fields, model)
        
        # Generate document content using AI
        if self.ai_manager:
            content = self._ai_generate_content(document_type, data, template, model)
        else:
            content = self._template_generate_content(document_type, data, template)
        
        return {
            "document_type": document_type,
            "generated_date": datetime.now().isoformat(),
            "data": data,
            "content": content,
            "template_info": template
        }
    
    def _ai_generate_content(
        self, 
        document_type: str, 
        data: Dict[str, Any], 
        template: Dict[str, Any],
        model: str
    ) -> str:
        """Use AI to generate comprehensive document content."""
        
        prompt = f"""
Generate a comprehensive {document_type} for a medical device company following ISO13485 requirements.

Document Structure Required:
{json.dumps(template["sections"], indent=2)}

Input Data:
{json.dumps(data, indent=2, default=str)}

Requirements:
1. Follow ISO13485:2016 standards
2. Include all required sections with appropriate detail
3. Use professional medical device industry language
4. Ensure regulatory compliance
5. Include specific field requirements and validation criteria
6. Add appropriate references to ISO13485 clauses where relevant
7. Format as structured markdown with clear headings
8. Include placeholder fields for signatures and dates where appropriate

Generate the complete document content now:
"""
        
        messages = [
            {
                "role": "system", 
                "content": """You are an expert ISO13485 documentation specialist. Generate professional, 
                compliant medical device quality documents that meet regulatory standards. Ensure all 
                content is accurate, complete, and follows industry best practices."""
            },
            {"role": "user", "content": prompt}
        ]
        
        try:
            provider, model_name = self.ai_manager.select_best_model("document_generation", model)
            response = self.ai_manager.generate_completion(
                messages=messages,
                provider=provider,
                model=model_name,
                temperature=0.3,  # Lower temperature for consistent, professional output
                max_tokens=4000
            )
            return response.content
            
        except Exception as e:
            st.warning(f"AI generation failed, using template: {str(e)}")
            return self._template_generate_content(document_type, data, template)
    
    def _ai_fill_missing_fields(
        self, 
        document_type: str, 
        data: Dict[str, Any], 
        missing_fields: List[str],
        model: str
    ) -> Dict[str, Any]:
        """Use AI to suggest content for missing required fields."""
        
        prompt = f"""
For a {document_type}, suggest appropriate content for these missing fields based on the provided context:

Missing Fields: {missing_fields}
Existing Data: {json.dumps(data, indent=2, default=str)}

Provide suggestions as JSON with the missing field names as keys and appropriate professional content as values.
Make suggestions specific to medical device quality management and ISO13485 compliance.
"""
        
        messages = [
            {"role": "system", "content": "You are an ISO13485 documentation expert."},
            {"role": "user", "content": prompt}
        ]
        
        try:
            provider, model_name = self.ai_manager.select_best_model("document_generation", model)
            response = self.ai_manager.generate_completion(
                messages=messages,
                provider=provider,
                model=model_name,
                temperature=0.4
            )
            
            suggestions = json.loads(response.content)
            
            # Merge suggestions into data
            for field, suggestion in suggestions.items():
                if field in missing_fields and not data.get(field):
                    data[field] = f"[AI Suggestion] {suggestion}"
            
            return data
            
        except Exception as e:
            # Return original data if AI suggestion fails
            return data
    
    def _template_generate_content(
        self, 
        document_type: str, 
        data: Dict[str, Any], 
        template: Dict[str, Any]
    ) -> str:
        """Generate document content using templates when AI is not available."""
        
        templates = {
            "CAPA Request Form": self._generate_capa_template,
            "Nonconformance Report": self._generate_ncr_template,
            "Internal Audit Checklist": self._generate_audit_template,
            "Risk Assessment Form": self._generate_risk_template,
            "Management Review Agenda": self._generate_mgmt_review_template,
            "Supplier Audit Report": self._generate_supplier_audit_template
        }
        
        generator = templates.get(document_type, self._generate_generic_template)
        return generator(data, template)
    
    def _generate_capa_template(self, data: Dict[str, Any], template: Dict[str, Any]) -> str:
        """Generate CAPA document template."""
        return f"""
# CORRECTIVE AND PREVENTIVE ACTION (CAPA) REQUEST

## Document Information
- **CAPA Number:** {data.get('capa_number', 'TBD')}
- **Date Initiated:** {data.get('generated_date', datetime.now().strftime('%Y-%m-%d'))}
- **Prepared By:** {data.get('assigned_to', 'TBD')}
- **Department:** {data.get('department', 'Quality')}

## Product Information
- **Product Name:** {data.get('product_name', 'TBD')}
- **Product ID/SKU:** {data.get('product_id', 'TBD')}
- **Priority Level:** {data.get('priority', 'Medium')}

## Problem Description
**Issue Description:**
{data.get('issue_description', 'Issue description to be provided')}

**Impact Assessment:**
- Severity: {data.get('priority', 'Medium')}
- Affected Products: {data.get('product_name', 'TBD')}
- Customer Impact: To be assessed

## Root Cause Analysis
**Root Cause:**
{data.get('root_cause', 'Root cause analysis to be completed')}

**Analysis Method:** 5 Whys, Fishbone Diagram, or equivalent systematic approach

## Corrective Actions
**Immediate Actions Taken:**
1. Quarantine affected inventory
2. Notify relevant stakeholders
3. Document incident

**Corrective Action Plan:**
{data.get('corrective_action', 'Corrective actions to be determined')}

**Implementation Timeline:** {data.get('due_date', 'TBD')}
**Responsible Person:** {data.get('assigned_to', 'TBD')}

## Preventive Actions
**Preventive Action Plan:**
{data.get('preventive_action', 'Preventive actions to be determined')}

**Scope of Implementation:** System-wide review and improvement

## Effectiveness Verification
**Verification Method:** 
- Monitoring and measurement
- Internal audit verification
- Management review

**Success Criteria:** Elimination of root cause and prevention of recurrence

**Verification Timeline:** 90 days post-implementation

## Approval
**Quality Manager Approval:**
Name: _________________ Signature: _________________ Date: _______

**Management Representative Approval:** 
Name: _________________ Signature: _________________ Date: _______

---
*This document is controlled per ISO13485:2016 requirements*
"""
    
    def _generate_ncr_template(self, data: Dict[str, Any], template: Dict[str, Any]) -> str:
        """Generate Nonconformance Report template."""
        return f"""
# NONCONFORMANCE REPORT (NCR)

## Document Information
- **NCR Number:** {data.get('ncr_number', 'NCR-' + datetime.now().strftime('%Y%m%d-001'))}
- **Date Initiated:** {datetime.now().strftime('%Y-%m-%d')}
- **Initiated By:** {data.get('assigned_to', 'TBD')}
- **Department:** {data.get('department', 'Quality')}

## Product/Process Information
- **Product Name:** {data.get('product_name', 'TBD')}
- **Product ID:** {data.get('product_id', 'TBD')}
- **Lot/Serial Number:** {data.get('lot_number', 'TBD')}
- **Quantity Affected:** {data.get('quantity', 'TBD')}

## Nonconformance Description
**Description of Nonconformance:**
{data.get('nonconformance_description', data.get('issue_description', 'Nonconformance description to be provided'))}

**Detection Method:**
☐ Incoming Inspection  ☐ In-Process Inspection  ☐ Final Inspection
☐ Customer Complaint   ☐ Internal Audit         ☐ Other: ________

## Impact Assessment
**Severity Level:** {data.get('priority', 'Medium')}
**Customer Impact:** To be evaluated
**Regulatory Impact:** To be assessed

## Disposition
☐ Accept As Is        ☐ Rework              ☐ Return to Supplier
☐ Scrap              ☐ Use Alternative     ☐ Other: ________

**Justification:** {data.get('disposition_justification', 'To be provided')}

## Investigation
**Investigation Required:** ☐ Yes  ☐ No

**Investigation Findings:**
{data.get('investigation_findings', 'Investigation to be completed if required')}

## Corrective Action
**Corrective Action Required:** ☐ Yes  ☐ No

If Yes, CAPA Number: ________________

## Closure
**Closed By:** _________________
**Date:** _______
**Signature:** _________________

---
*Document controlled per ISO13485:2016 Section 4.2*
"""
    
    def _generate_audit_template(self, data: Dict[str, Any], template: Dict[str, Any]) -> str:
        """Generate Internal Audit Checklist template."""
        return f"""
# INTERNAL AUDIT CHECKLIST - ISO13485:2016

## Audit Information
- **Audit Number:** {data.get('audit_number', 'IA-' + datetime.now().strftime('%Y-%m-%d'))}
- **Audit Date:** {data.get('audit_date', datetime.now().strftime('%Y-%m-%d'))}
- **Lead Auditor:** {data.get('auditor_name', 'TBD')}
- **Audit Team:** {data.get('audit_team', 'TBD')}

## Audit Scope
**Areas to be Audited:**
{data.get('audit_scope', 'Quality Management System elements per ISO13485:2016')}

**Departments/Processes:**
- Quality Management
- Design and Development
- Manufacturing
- Supplier Management

## ISO13485:2016 Requirements Checklist

### 4. Quality Management System
- [ ] 4.1 General requirements
- [ ] 4.2 Documentation requirements

### 5. Management Responsibility  
- [ ] 5.1 Management commitment
- [ ] 5.2 Customer focus
- [ ] 5.3 Quality policy
- [ ] 5.4 Planning
- [ ] 5.5 Responsibility, authority and communication
- [ ] 5.6 Management review

### 6. Resource Management
- [ ] 6.1 Provision of resources
- [ ] 6.2 Human resources
- [ ] 6.3 Infrastructure
- [ ] 6.4 Work environment and contamination control

### 7. Product Realization
- [ ] 7.1 Planning of product realization
- [ ] 7.2 Customer-related processes
- [ ] 7.3 Design and development
- [ ] 7.4 Purchasing
- [ ] 7.5 Production and service provision
- [ ] 7.6 Control of monitoring and measuring equipment

### 8. Measurement, Analysis and Improvement
- [ ] 8.1 General
- [ ] 8.2 Monitoring and measurement
- [ ] 8.3 Control of nonconforming product
- [ ] 8.4 Analysis of data
- [ ] 8.5 Improvement

## Findings Summary
**Conformances:** ____
**Minor Nonconformances:** ____
**Major Nonconformances:** ____
**Observations:** ____

## Audit Conclusion
**Overall Assessment:** System appears to be effectively implemented ☐ / requires improvement ☐

**Auditor Signature:** _______________ **Date:** _______

---
*Internal audit conducted per ISO13485:2016 Section 8.2.2*
"""
    
    def _generate_generic_template(self, data: Dict[str, Any], template: Dict[str, Any]) -> str:
        """Generate generic template for unsupported document types."""
        content = f"# {data.get('document_type', 'ISO13485 Document')}\n\n"
        content += f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        
        for section in template.get("sections", []):
            content += f"## {section}\n\n"
            content += "[Content to be added]\n\n"
        
        return content
    
    def export_to_docx(self, generated_doc: Dict[str, Any]) -> BytesIO:
        """Export generated document to Word format with professional formatting."""
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # Add document title
        title = doc.add_heading(generated_doc["document_type"], 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_format = title.runs[0].font
        title_format.name = self.formatting["title_font"]
        title_format.size = Pt(self.formatting["title_size"])
        title_format.color.rgb = self.formatting["colors"]["primary"]
        
        # Add generation info
        info_para = doc.add_paragraph()
        info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        info_run = info_para.add_run(f"Generated: {generated_doc['generated_date']}")
        info_run.font.size = Pt(10)
        info_run.font.color.rgb = self.formatting["colors"]["secondary"]
        
        doc.add_paragraph()  # Spacing
        
        # Process markdown content into Word
        content_lines = generated_doc["content"].split('\n')
        current_paragraph = None
        
        for line in content_lines:
            line = line.strip()
            
            if not line:
                if current_paragraph:
                    doc.add_paragraph()
                continue
            
            # Handle headers
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                header_text = line.lstrip('#').strip()
                
                if level == 1:
                    header = doc.add_heading(header_text, 1)
                    header_format = header.runs[0].font
                    header_format.name = self.formatting["header_font"]
                    header_format.color.rgb = self.formatting["colors"]["primary"]
                elif level == 2:
                    header = doc.add_heading(header_text, 2)
                    header_format = header.runs[0].font
                    header_format.name = self.formatting["header_font"]
                    header_format.color.rgb = self.formatting["colors"]["secondary"]
                else:
                    header = doc.add_heading(header_text, 3)
                    header_format = header.runs[0].font
                    header_format.name = self.formatting["header_font"]
                
                current_paragraph = None
                continue
            
            # Handle bullet points
            if line.startswith('- ') or line.startswith('* '):
                bullet_text = line[2:].strip()
                para = doc.add_paragraph(bullet_text, style='List Bullet')
                para.runs[0].font.name = self.formatting["body_font"]
                para.runs[0].font.size = Pt(self.formatting["body_size"])
                current_paragraph = None
                continue
            
            # Handle numbered lists
            if line[0].isdigit() and line[1:3] == '. ':
                list_text = line[3:].strip()
                para = doc.add_paragraph(list_text, style='List Number')
                para.runs[0].font.name = self.formatting["body_font"]
                para.runs[0].font.size = Pt(self.formatting["body_size"])
                current_paragraph = None
                continue
            
            # Handle checkboxes
            if '☐' in line or '☑' in line:
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.font.name = self.formatting["body_font"]
                run.font.size = Pt(self.formatting["body_size"])
                current_paragraph = None
                continue
            
            # Handle bold text (simple **text** pattern)
            if '**' in line:
                para = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    run = para.add_run(part)
                    run.font.name = self.formatting["body_font"]
                    run.font.size = Pt(self.formatting["body_size"])
                    if i % 2 == 1:  # Odd indices are bold text
                        run.font.bold = True
                current_paragraph = None
                continue
            
            # Regular text
            if current_paragraph is None:
                current_paragraph = doc.add_paragraph()
            
            run = current_paragraph.add_run(line + ' ')
            run.font.name = self.formatting["body_font"]
            run.font.size = Pt(self.formatting["body_size"])
        
        # Add footer
        doc.add_page_break()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_run = footer_para.add_run("This document is controlled per ISO13485:2016 requirements")
        footer_run.font.size = Pt(9)
        footer_run.font.italic = True
        footer_run.font.color.rgb = self.formatting["colors"]["secondary"]
        
        # Save to BytesIO
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer
    
    def get_available_templates(self) -> List[str]:
        """Get list of available document templates."""
        return list(self.document_templates.keys())
    
    def validate_document_data(self, document_type: str, data: Dict[str, Any]) -> Dict[str, List[str]]:
        """Validate document data against template requirements."""
        if document_type not in self.document_templates:
            return {"errors": [f"Unsupported document type: {document_type}"]}
        
        template = self.document_templates[document_type]
        errors = []
        warnings = []
        
        # Check required fields
        for field in template["required_fields"]:
            if not data.get(field):
                errors.append(f"Required field missing: {field}")
        
        # Check data types and formats
        if data.get('due_date'):
            try:
                datetime.strptime(str(data['due_date']), '%Y-%m-%d')
            except ValueError:
                warnings.append("Due date format should be YYYY-MM-DD")
        
        return {"errors": errors, "warnings": warnings}
